import requests
import os
from docx import Document

def create_dummy_cv(filename="test_cv.docx"):
    doc = Document()
    doc.add_heading('Ali Veli', 0)
    
    doc.add_paragraph('Email: ali.veli@example.com | Tel: 0555 123 45 67')
    doc.add_paragraph('GitHub: https://github.com/aliveli')
    
    doc.add_heading('Deneyim', level=1)
    doc.add_paragraph('Yazılım Geliştirici - ABC Teknoloji')
    doc.add_paragraph('Python, FastAPI ve React kullanarak projeler geliştirdim.')
    
    doc.add_heading('Yetenekler', level=1)
    doc.add_paragraph('Python, Java, Docker, AWS, React, SQL')
    
    doc.add_heading('Eğitim', level=1)
    doc.add_paragraph('Bilgisayar Mühendisliği - XYZ Üniversitesi')
    
    doc.save(filename)
    return filename

def test_upload():
    filename = create_dummy_cv()
    url = "http://127.0.0.1:8000/api/v1/cv/upload"
    
    print(f"Test dosyası oluşturuldu: {filename}")
    print(f"API'ye gönderiliyor: {url}")
    
    try:
        with open(filename, 'rb') as f:
            files = {'file': (filename, f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            data = {'job_description': "Aranan Nitelikler: Python ve React konusunda deneyimli, Docker ve AWS bilen yazılım geliştirici aranıyor."}
            response = requests.post(url, files=files, data=data)
            
        if response.status_code == 200:
            print("\n✅ Başarılı! Sunucu yanıtı:")
            import json
            print(json.dumps(response.json(), indent=2, ensure_ascii=False))
        else:
            print(f"\n❌ Hata! Durum Kodu: {response.status_code}")
            print(response.text)
            
    except Exception as e:
        print(f"\n❌ Bağlantı Hatası: {str(e)}")
    finally:
        if os.path.exists(filename):
            os.remove(filename)
            print(f"\nTest dosyası temizlendi: {filename}")

if __name__ == "__main__":
    test_upload()
